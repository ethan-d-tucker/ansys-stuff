"""
MIL-STD-810H tailored DOCX report generator for composite PSD analysis.

Generates a professional report referencing Method 514.8, with:
- Requirements traceability (REQ-VIB-001 through 004)
- Multi-environment PSD results (per axis)
- Composite failure assessment (Tsai-Wu, Max Stress, FoS)
- Overall compliance matrix and pass/fail determination

Reuses styling patterns from run_and_report.py (python-docx).
"""

import os
import datetime

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    import pyvista as pv
    HAS_PYVISTA = True
except ImportError:
    HAS_PYVISTA = False


# ---------------------------------------------------------------------------
# Colour theme
# ---------------------------------------------------------------------------
HEADING_COLOR = RGBColor(0x15, 0x65, 0xC0)
PASS_COLOR = RGBColor(0x2E, 0x7D, 0x32)  # green
FAIL_COLOR = RGBColor(0xC6, 0x28, 0x28)  # red
CAPTION_COLOR = RGBColor(0x60, 0x7D, 0x8B)


# ---------------------------------------------------------------------------
# DOCX helper functions  (same pattern as run_and_report.py)
# ---------------------------------------------------------------------------

def _setup_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    return doc


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = HEADING_COLOR
    return h


def _add_figure(doc, path, caption, width=Inches(6.0)):
    if path and os.path.isfile(path):
        doc.add_picture(path, width=width)
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = CAPTION_COLOR
        return True
    return False


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers),
                          style="Light Grid Accent 1")
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def _pass_fail_text(passed):
    return "PASS" if passed else "FAIL"


# ---------------------------------------------------------------------------
# Matplotlib chart generators
# ---------------------------------------------------------------------------

def _save_psd_overlay(profiles, modal_freqs, output_path):
    """Plot all PSD profiles overlaid on one log-log chart."""
    fig, ax = plt.subplots(figsize=(8, 5))
    colors = ["#1565C0", "#E65100", "#2E7D32", "#6A1B9A"]

    for i, prof in enumerate(profiles):
        freqs = [f for f, _ in prof["psd_table"]]
        vals = [v for _, v in prof["psd_table"]]
        c = colors[i % len(colors)]
        ax.loglog(freqs, vals, "o-", color=c, linewidth=2, markersize=5,
                  label=f'{prof["requirement_id"]}: {prof["name"]}')

    for j, f in enumerate(modal_freqs[:6]):
        ax.axvline(f, color="red", linestyle="--", alpha=0.4, linewidth=0.8)
        if j < 4:
            ax.text(f * 1.05, ax.get_ylim()[1] * 0.7, f"f{j+1}={f:.0f}",
                    fontsize=7, color="red", rotation=90, va="top")

    ax.set_xlabel("Frequency (Hz)")
    ax.set_ylabel("PSD (G^2/Hz)")
    ax.set_title("MIL-STD-810H Vibration Environment PSD Profiles")
    ax.legend(fontsize=8, loc="lower left")
    ax.grid(True, which="both", alpha=0.3)
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return output_path


def _save_fos_chart(env_results, failure_data, required_fos, output_path):
    """Horizontal bar chart of FoS per environment with threshold lines."""
    labels = []
    fos_tw = []
    fos_ms = []

    for pid, axes_data in env_results.items():
        for axis, psd_res in axes_data.items():
            key = (pid, axis)
            if key in failure_data:
                fr = failure_data[key]
                labels.append(f"{pid}\n{axis}-axis")
                fos_tw.append(min(fr.min_fos_tw, 20.0))
                fos_ms.append(min(fr.min_fos_ms, 20.0))

    if not labels:
        return None

    y_pos = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(9, max(4, len(labels) * 0.6)))
    bar_h = 0.35
    ax.barh(y_pos - bar_h / 2, fos_tw, bar_h, label="Tsai-Wu FoS",
            color="#1565C0", alpha=0.85)
    ax.barh(y_pos + bar_h / 2, fos_ms, bar_h, label="Max Stress FoS",
            color="#E65100", alpha=0.85)

    ax.axvline(1.0, color="red", linewidth=2, linestyle="-", label="FoS = 1.0 (Failure)")
    ax.axvline(required_fos, color="orange", linewidth=2, linestyle="--",
               label=f"FoS = {required_fos} (Required)")

    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("Factor of Safety")
    ax.set_title("Composite Factor of Safety -- All Environments")
    ax.legend(fontsize=8, loc="lower right")
    ax.grid(True, axis="x", alpha=0.3)
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return output_path


def _save_failure_bar_chart(env_results, failure_data, output_path):
    """Grouped bar chart of max failure index per environment/axis."""
    labels = []
    tw_vals = []
    ms_vals = []

    for pid, axes_data in env_results.items():
        for axis, psd_res in axes_data.items():
            key = (pid, axis)
            if key in failure_data:
                fr = failure_data[key]
                labels.append(f"{pid}\n{axis}")
                tw_vals.append(fr.max_tw_index)
                ms_vals.append(fr.max_ms_index)

    if not labels:
        return None

    x = np.arange(len(labels))
    w = 0.35
    fig, ax = plt.subplots(figsize=(max(8, len(labels) * 1.2), 5))
    ax.bar(x - w / 2, tw_vals, w, label="Tsai-Wu Index", color="#1565C0")
    ax.bar(x + w / 2, ms_vals, w, label="Max Stress Index", color="#E65100")
    ax.axhline(1.0, color="red", linewidth=2, linestyle="--", label="Failure Threshold")
    ax.set_xticks(x)
    ax.set_xticklabels(labels, fontsize=8)
    ax.set_ylabel("Failure Index")
    ax.set_title("Composite Failure Indices by Environment and Axis")
    ax.legend(fontsize=8)
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return output_path


def _save_layup_chart(layup, output_path):
    """Stacked ply visualization."""
    fig, ax = plt.subplots(figsize=(6, 4))
    y = 0.0
    colors_map = {"carbon_epoxy_woven": "#37474F", "honeycomb_core": "#FFB300"}
    labels_seen = set()

    for ply in layup:
        t = ply["thickness_mm"]
        mat = ply["mat"]
        c = colors_map.get(mat, "#90A4AE")
        label = mat.replace("_", " ").title() if mat not in labels_seen else None
        labels_seen.add(mat)
        ax.barh(y + t / 2, 1.0, t, color=c, edgecolor="black", linewidth=0.5,
                label=label)
        ax.text(0.5, y + t / 2, f'{ply["angle"]}deg  {t:.2f}mm',
                ha="center", va="center", fontsize=7, color="white" if "carbon" in mat else "black")
        y += t

    ax.set_xlim(0, 1)
    ax.set_ylim(0, y)
    ax.set_ylabel("Thickness (mm)")
    ax.set_xticks([])
    ax.set_title("Composite Layup Cross-Section")
    ax.legend(fontsize=8, loc="upper right")
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return output_path


def _save_freq_bar_chart(nat_freqs, output_path):
    """Bar chart of natural frequencies."""
    n = len(nat_freqs)
    fig, ax = plt.subplots(figsize=(max(6, n * 0.5), 4))
    colors = plt.cm.viridis(np.linspace(0.2, 0.8, n))
    bars = ax.bar(range(1, n + 1), nat_freqs, color=colors)
    for bar, f in zip(bars, nat_freqs):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 10,
                f"{f:.0f}", ha="center", va="bottom", fontsize=7)
    ax.set_xlabel("Mode Number")
    ax.set_ylabel("Frequency (Hz)")
    ax.set_title("Natural Frequencies")
    ax.grid(True, axis="y", alpha=0.3)
    fig.tight_layout()
    fig.savefig(output_path, dpi=150)
    plt.close(fig)
    return output_path


# ---------------------------------------------------------------------------
# PyVista contour image capture
# ---------------------------------------------------------------------------

def _capture_contour(grid, scalars, title, output_path, cmap="jet"):
    """Render a scalar contour on the mesh and save as PNG."""
    if not HAS_PYVISTA or grid is None:
        return None
    try:
        pv.OFF_SCREEN = True
        pl = pv.Plotter(off_screen=True, window_size=(1920, 1080))
        g = grid.copy()
        g.point_data[title] = scalars
        pl.add_mesh(g, scalars=title, cmap=cmap, show_edges=False,
                     scalar_bar_args={"title": title})
        pl.view_isometric()
        pl.screenshot(output_path)
        pl.close()
        return output_path
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Main report builder
# ---------------------------------------------------------------------------

def generate_milstd_report(
    modal_data,
    env_results,          # {profile_id: {axis: PSDResults}}
    failure_data,         # {(profile_id, axis): FailureResult}
    core_failure_data,    # {(profile_id, axis): dict} or {}
    profiles,             # list[dict] from mil_std_profiles
    layup,                # list[dict] from material_library
    material_info,        # dict with elastic + strength
    config,               # SimulationConfig
    output_dir,
    required_fos=1.5,
):
    """
    Generate the full MIL-STD-810H DOCX report.

    Returns path to the generated DOCX file.
    """
    os.makedirs(output_dir, exist_ok=True)
    img_dir = os.path.join(output_dir, "images")
    os.makedirs(img_dir, exist_ok=True)

    doc = _setup_doc()
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

    # ======= Generate charts =======
    print("  Generating charts ...")
    chart_paths = {}

    chart_paths["psd_overlay"] = _save_psd_overlay(
        profiles, modal_data.nat_freqs,
        os.path.join(img_dir, "psd_overlay.png"),
    )
    chart_paths["fos"] = _save_fos_chart(
        env_results, failure_data, required_fos,
        os.path.join(img_dir, "fos_chart.png"),
    )
    chart_paths["failure_bar"] = _save_failure_bar_chart(
        env_results, failure_data,
        os.path.join(img_dir, "failure_indices.png"),
    )
    chart_paths["layup"] = _save_layup_chart(
        layup, os.path.join(img_dir, "layup.png"),
    )
    chart_paths["freq_bar"] = _save_freq_bar_chart(
        modal_data.nat_freqs,
        os.path.join(img_dir, "frequencies.png"),
    )

    # Capture contour images for worst-case environment
    worst_pid, worst_axis, worst_fr = _find_worst_case(failure_data)
    contour_paths = {}
    if worst_pid and modal_data.grid_base is not None:
        psd_worst = env_results[worst_pid][worst_axis]
        contour_paths["stress_eqv"] = _capture_contour(
            modal_data.grid_base, psd_worst.stress_eqv,
            "1-sigma von Mises (Pa)",
            os.path.join(img_dir, "contour_stress_eqv.png"),
        )
        contour_paths["disp_mag"] = _capture_contour(
            modal_data.grid_base, psd_worst.disp_mag,
            "1-sigma |U| (m)",
            os.path.join(img_dir, "contour_disp_mag.png"),
        )
        if worst_fr is not None:
            contour_paths["tsai_wu"] = _capture_contour(
                modal_data.grid_base, worst_fr.tsai_wu_index,
                "Tsai-Wu Failure Index",
                os.path.join(img_dir, "contour_tsai_wu.png"),
                cmap="RdYlGn_r",
            )
            contour_paths["fos_tw"] = _capture_contour(
                modal_data.grid_base,
                np.minimum(worst_fr.tsai_wu_fos, 20.0),
                "Tsai-Wu Factor of Safety",
                os.path.join(img_dir, "contour_fos_tw.png"),
                cmap="RdYlGn",
            )

    # ======================================================================
    # TITLE PAGE
    # ======================================================================
    doc.add_paragraph("")
    title = doc.add_heading(
        "MIL-STD-810H Random Vibration (PSD)\nAnalysis Report", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = HEADING_COLOR

    sub = doc.add_paragraph(
        f"{config.part_name} -- Composite Construction"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)

    doc.add_paragraph("")
    meta_rows = [
        ("Date", now),
        ("Standard", "MIL-STD-810H, Method 514.8"),
        ("Software", "ANSYS Mechanical APDL (Student 2025 R2, v25.2)"),
        ("Geometry", os.path.basename(config.geometry_file)),
        ("Element Type", "SOLID187 (10-node tetrahedral)"),
    ]
    _add_table(doc, ["Parameter", "Value"], meta_rows)
    doc.add_page_break()

    # ======================================================================
    # 1. EXECUTIVE SUMMARY
    # ======================================================================
    _add_heading(doc, "1. Executive Summary")

    # Overall pass/fail
    all_pass_tw = all(fr.overall_pass_tw for fr in failure_data.values())
    all_pass_ms = all(fr.overall_pass_ms for fr in failure_data.values())
    min_fos_tw = min(fr.min_fos_tw for fr in failure_data.values()) if failure_data else 0
    min_fos_ms = min(fr.min_fos_ms for fr in failure_data.values()) if failure_data else 0
    meets_required = min_fos_tw >= required_fos and min_fos_ms >= required_fos

    verdict = "PASS" if meets_required else "FAIL"
    doc.add_paragraph(
        f"This report documents the random vibration (PSD) analysis of the "
        f"{config.part_name} composite structure per MIL-STD-810H Method 514.8.  "
        f"{len(profiles)} vibration environments were evaluated across three "
        f"orthogonal excitation axes."
    )

    p = doc.add_paragraph()
    run = p.add_run(f"Overall Qualification Result: {verdict}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = PASS_COLOR if meets_required else FAIL_COLOR

    _add_heading(doc, "1.1 Key Results Summary", level=2)
    summary_rows = []
    for prof in profiles:
        pid = prof["id"]
        for axis in ["X", "Y", "Z"]:
            if pid in env_results and axis in env_results[pid]:
                psd_res = env_results[pid][axis]
                key = (pid, axis)
                fr = failure_data.get(key)
                if fr:
                    summary_rows.append([
                        prof["requirement_id"],
                        prof["name"][:25],
                        axis,
                        f"{psd_res.max_disp_mag_um:.2f}",
                        f"{psd_res.max_stress_eqv_mpa:.4f}",
                        f"{fr.max_tw_index:.4f}",
                        f"{fr.min_fos_tw:.1f}",
                        _pass_fail_text(fr.min_fos_tw >= required_fos),
                    ])

    _add_table(doc,
               ["Req ID", "Environment", "Axis", "Max |U| (um)",
                "Max SEQV (MPa)", "TW Index", "FoS_TW", "Result"],
               summary_rows)

    doc.add_paragraph("")
    if contour_paths.get("stress_eqv"):
        _add_figure(doc, contour_paths["stress_eqv"],
                    f"Figure 1-1: Worst-case 1-sigma von Mises stress "
                    f"({worst_pid}, {worst_axis}-axis)")

    doc.add_page_break()

    # ======================================================================
    # 2. APPLICABLE DOCUMENTS AND REQUIREMENTS
    # ======================================================================
    _add_heading(doc, "2. Applicable Documents and Requirements")

    _add_heading(doc, "2.1 MIL-STD-810H Method 514.8 Overview", level=2)
    doc.add_paragraph(
        "MIL-STD-810H Method 514.8 defines environmental vibration test procedures "
        "and tailoring guidelines for materiel intended for use in military and "
        "commercial applications.  The method covers both sinusoidal and random "
        "vibration, with categories representing different operational platforms "
        "(ground vehicles, rotary-wing aircraft, fixed-wing aircraft, etc.)."
    )
    doc.add_paragraph(
        "This analysis evaluates the structural response under random vibration "
        "power spectral density (PSD) excitation corresponding to the selected "
        "environmental categories.  The assessment includes composite failure "
        "analysis using industry-standard criteria (Tsai-Wu and Maximum Stress) "
        "with applied factors of safety."
    )

    _add_heading(doc, "2.2 Requirements Traceability", level=2)
    req_rows = []
    for prof in profiles:
        req_rows.append([
            prof["requirement_id"],
            prof["name"],
            prof["mil_std_ref"],
            f"{prof['grms']:.2f} Grms",
            f"{prof['duration_min_per_axis']} min/axis",
        ])
    _add_table(doc,
               ["Req ID", "Environment", "MIL-STD Reference", "Grms", "Duration"],
               req_rows)

    _add_heading(doc, "2.3 Acceptance Criteria", level=2)
    doc.add_paragraph(
        f"Structural adequacy is assessed against the following criteria:"
    )
    doc.add_paragraph(
        f"  - Tsai-Wu failure index < 1.0 (no predicted failure)", style="List Bullet"
    )
    doc.add_paragraph(
        f"  - Maximum Stress failure index < 1.0", style="List Bullet"
    )
    doc.add_paragraph(
        f"  - Factor of Safety (FoS) >= {required_fos:.1f} "
        f"for all environments and axes", style="List Bullet"
    )

    doc.add_page_break()

    # ======================================================================
    # 3. MODEL DESCRIPTION
    # ======================================================================
    _add_heading(doc, "3. Model Description")

    _add_heading(doc, "3.1 Part Geometry", level=2)
    mi = modal_data.mesh_info
    doc.add_paragraph(
        f"Part: {config.part_name}. "
        f"Geometry imported from {os.path.basename(config.geometry_file)} "
        f"via Parasolid ac4 converter."
    )
    geo = mi.get("geometry", {})
    if geo:
        doc.add_paragraph(
            f"Volumes: {geo.get('n_volu', 'N/A')}, Areas: {geo.get('n_area', 'N/A')}"
        )

    dims_rows = []
    for i, ax in enumerate(["X", "Y", "Z"]):
        dims_rows.append([
            ax,
            f"{mi['mins_m'][i]*1e3:.1f}",
            f"{mi['maxs_m'][i]*1e3:.1f}",
            f"{mi['spans_m'][i]*1e3:.1f}",
        ])
    _add_table(doc, ["Axis", "Min (mm)", "Max (mm)", "Span (mm)"], dims_rows)

    _add_heading(doc, "3.2 Finite Element Model", level=2)
    doc.add_paragraph(
        f"Element type: SOLID187 (10-node tetrahedral, quadratic).  "
        f"Free mesh with element size {config.element_size*1000:.1f} mm.  "
        f"Total mesh: {mi['n_nodes']} nodes, {mi['n_elements']} elements."
    )
    doc.add_paragraph(
        "Note: Homogenised orthotropic material properties are assigned to the "
        "solid elements.  Ply-level stress is approximated by rotating the "
        "homogenised stress tensor into each ply's material coordinate system."
    )

    _add_heading(doc, "3.3 Material Properties", level=2)
    elastic = config.material_props
    mat_rows = [
        ("EX (GPa)", f"{elastic.get('EX', 0)/1e9:.1f}"),
        ("EY (GPa)", f"{elastic.get('EY', 0)/1e9:.1f}"),
        ("EZ (GPa)", f"{elastic.get('EZ', 0)/1e9:.1f}"),
        ("GXY (GPa)", f"{elastic.get('GXY', 0)/1e9:.1f}"),
        ("GXZ (GPa)", f"{elastic.get('GXZ', 0)/1e9:.1f}"),
        ("GYZ (GPa)", f"{elastic.get('GYZ', 0)/1e9:.1f}"),
        ("PRXY", f"{elastic.get('PRXY', 0):.3f}"),
        ("PRXZ", f"{elastic.get('PRXZ', 0):.3f}"),
        ("PRYZ", f"{elastic.get('PRYZ', 0):.3f}"),
        ("Density (kg/m3)", f"{elastic.get('DENS', 0):.0f}"),
    ]
    _add_table(doc, ["Property", "Value"], mat_rows)

    _add_heading(doc, "3.4 Material Strength Allowables", level=2)
    doc.add_paragraph(
        "Strength allowables used for composite failure analysis.  "
        "Values are typical / B-basis for the material system."
    )
    strength = material_info.get("strength", {})
    str_rows = [
        ("Xt -- Tensile, fibre dir (MPa)", f"{strength.get('Xt', 0)/1e6:.0f}"),
        ("Xc -- Compressive, fibre dir (MPa)", f"{strength.get('Xc', 0)/1e6:.0f}"),
        ("Yt -- Tensile, transverse (MPa)", f"{strength.get('Yt', 0)/1e6:.0f}"),
        ("Yc -- Compressive, transverse (MPa)", f"{strength.get('Yc', 0)/1e6:.0f}"),
        ("S12 -- In-plane shear (MPa)", f"{strength.get('S12', 0)/1e6:.0f}"),
        ("S13 -- Interlaminar shear XZ (MPa)", f"{strength.get('S13', 0)/1e6:.0f}"),
        ("S23 -- Interlaminar shear YZ (MPa)", f"{strength.get('S23', 0)/1e6:.0f}"),
        ("Zt -- Through-thickness tensile (MPa)", f"{strength.get('Zt', 0)/1e6:.0f}"),
        ("Zc -- Through-thickness compressive (MPa)", f"{strength.get('Zc', 0)/1e6:.0f}"),
    ]
    _add_table(doc, ["Allowable", "Value"], str_rows)

    _add_heading(doc, "3.5 Composite Layup", level=2)
    from material_library import get_layup_summary
    ls = get_layup_summary(layup)
    doc.add_paragraph(
        f"Sandwich construction with {ls['n_plies']} plies "
        f"({ls['n_face_plies']} face-sheet plies + {ls['n_core_plies']} core).  "
        f"Total laminate thickness: {ls['total_thickness_mm']:.3f} mm.  "
        f"Stacking sequence: [{ls['stacking_sequence']}].  "
        f"Symmetric: {'Yes' if ls['symmetric'] else 'No'}."
    )

    ply_rows = []
    for p in layup:
        mat_short = "Carbon/Epoxy" if "carbon" in p["mat"] else "Honeycomb"
        ply_rows.append([
            str(p.get("ply", "")),
            mat_short,
            f"{p['thickness_mm']:.3f}",
            f"{p['angle']}",
            p.get("role", ""),
        ])
    _add_table(doc, ["Ply", "Material", "Thickness (mm)", "Angle (deg)", "Role"],
               ply_rows)

    doc.add_paragraph("")
    _add_figure(doc, chart_paths.get("layup"),
                "Figure 3-1: Composite layup cross-section")

    _add_heading(doc, "3.6 Boundary Conditions", level=2)
    bc = modal_data.bc_info
    doc.add_paragraph(
        f"Fixed support (UX=UY=UZ=0) applied at the {bc['axis']}-min end of the model.  "
        f"{bc['n_fixed']} nodes constrained.  "
        f"Base excitation applied independently along each orthogonal axis."
    )

    doc.add_page_break()

    # ======================================================================
    # 4. ANALYSIS CONFIGURATION
    # ======================================================================
    _add_heading(doc, "4. Analysis Configuration")

    _add_heading(doc, "4.1 Modal Analysis", level=2)
    doc.add_paragraph(
        f"Block Lanczos (LANB) eigensolver.  {config.num_modes} modes requested "
        f"in {config.freq_start:.0f}-{config.freq_end:.0f} Hz range.  "
        f"Sparse direct solver.  All modes expanded for stress recovery."
    )

    _add_heading(doc, "4.2 PSD Method", level=2)
    doc.add_paragraph(
        "Due to ANSYS Student Edition limitations (PSD combination results not "
        "stored), 1-sigma response is computed manually via Square Root of the "
        "Sum of Squares (SRSS) of modal contributions:"
    )
    doc.add_paragraph(
        "For each mode i: sigma_i^2 = integral[H_i^2(f) * S_a(f) df] / (2*pi*f_i)^4"
    )
    doc.add_paragraph(
        "1-sigma at node j: sigma_j = sqrt(sum_i(phi_ji^2 * sigma_i^2))"
    )
    doc.add_paragraph(
        f"Modal damping ratio: {config.damping_ratio*100:.0f}% critical (constant)."
    )

    _add_heading(doc, "4.3 Composite Failure Criteria", level=2)
    doc.add_paragraph(
        "Two failure criteria are evaluated at every node for each ply orientation:"
    )

    doc.add_paragraph(
        "Tsai-Wu Criterion (interactive, accounts for tension/compression asymmetry):",
        style="List Bullet",
    )
    doc.add_paragraph(
        "F_TW = F1*s1 + F2*s2 + F11*s1^2 + F22*s2^2 + F66*t12^2 + 2*F12*s1*s2"
    )
    doc.add_paragraph(
        "Where F1 = 1/Xt - 1/Xc, F11 = 1/(Xt*Xc), F12 = -0.5*sqrt(F11*F22), etc.  "
        "Failure predicted when F_TW >= 1.0."
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Maximum Stress Criterion (non-interactive, conservative baseline):",
        style="List Bullet",
    )
    doc.add_paragraph(
        "R = max(s1/Xt or |s1|/Xc, s2/Yt or |s2|/Yc, |t12|/S12).  "
        "Failure predicted when R >= 1.0."
    )
    doc.add_paragraph(
        f"Required Factor of Safety: {required_fos:.1f} "
        f"(FoS = 1 / failure_index)."
    )

    _add_heading(doc, "4.4 Vibration Environments", level=2)
    doc.add_paragraph(
        f"{len(profiles)} MIL-STD-810H environments selected for evaluation:"
    )
    for prof in profiles:
        doc.add_paragraph(
            f"{prof['requirement_id']} -- {prof['name']} "
            f"({prof['grms']:.2f} Grms, "
            f"{prof['psd_table'][0][0]:.0f}-{prof['psd_table'][-1][0]:.0f} Hz)",
            style="List Bullet",
        )

    doc.add_paragraph("")
    _add_figure(doc, chart_paths.get("psd_overlay"),
                "Figure 4-1: PSD profiles for all vibration environments "
                "with modal frequency markers")

    doc.add_page_break()

    # ======================================================================
    # 5. MODAL ANALYSIS RESULTS
    # ======================================================================
    _add_heading(doc, "5. Modal Analysis Results")

    _add_heading(doc, "5.1 Natural Frequencies", level=2)
    freq_rows = []
    for i, f in enumerate(modal_data.nat_freqs, 1):
        freq_rows.append([str(i), f"{f:.2f}"])
    _add_table(doc, ["Mode", "Frequency (Hz)"], freq_rows)

    doc.add_paragraph("")
    _add_figure(doc, chart_paths.get("freq_bar"),
                "Figure 5-1: Natural frequencies")

    doc.add_page_break()

    # ======================================================================
    # 6. PSD RESPONSE RESULTS -- PER ENVIRONMENT
    # ======================================================================
    _add_heading(doc, "6. PSD Response Results")

    fig_num = 1
    for prof_idx, prof in enumerate(profiles):
        pid = prof["id"]
        if pid not in env_results:
            continue

        _add_heading(doc,
                     f"6.{prof_idx+1} {prof['name']} ({prof['requirement_id']})",
                     level=2)
        doc.add_paragraph(prof["description"])

        # PSD input table
        psd_rows = [[f"{f:.1f}", f"{v:.6f}"] for f, v in prof["psd_table"]]
        psd_rows.append(["Overall Grms", f"{prof['grms']:.2f}"])
        _add_table(doc, ["Frequency (Hz)", "PSD (G^2/Hz)"], psd_rows)

        # Per-axis results table
        axis_rows = []
        for axis in ["X", "Y", "Z"]:
            if axis not in env_results[pid]:
                continue
            psd_res = env_results[pid][axis]
            key = (pid, axis)
            fr = failure_data.get(key)
            fos_str = f"{fr.min_fos_tw:.1f}" if fr else "N/A"
            result_str = _pass_fail_text(fr.min_fos_tw >= required_fos) if fr else "N/A"
            axis_rows.append([
                f"{axis}-axis",
                f"{psd_res.max_disp_mag_um:.2f}",
                f"{psd_res.max_stress_eqv_mpa:.4f}",
                fos_str,
                result_str,
            ])

        doc.add_paragraph("")
        _add_table(doc,
                   ["Excitation", "Max |U| (um)", "Max SEQV (MPa)",
                    "Min FoS (TW)", "Result"],
                   axis_rows)

        # Capture per-environment contour for worst axis
        worst_axis_for_env = _worst_axis_for_env(pid, env_results, failure_data)
        if worst_axis_for_env and modal_data.grid_base is not None:
            psd_w = env_results[pid][worst_axis_for_env]
            img_path = os.path.join(img_dir, f"stress_{pid}_{worst_axis_for_env}.png")
            cpath = _capture_contour(
                modal_data.grid_base, psd_w.stress_eqv,
                f"SEQV {pid} {worst_axis_for_env}",
                img_path,
            )
            if cpath:
                doc.add_paragraph("")
                _add_figure(doc, cpath,
                            f"Figure 6-{fig_num}: 1-sigma von Mises stress -- "
                            f"{prof['name']}, {worst_axis_for_env}-axis excitation")
                fig_num += 1

        doc.add_paragraph("")

    doc.add_page_break()

    # ======================================================================
    # 7. COMPOSITE FAILURE ASSESSMENT
    # ======================================================================
    _add_heading(doc, "7. Composite Failure Assessment")

    _add_heading(doc, "7.1 Methodology", level=2)
    doc.add_paragraph(
        "The global 1-sigma stress tensor (SX, SY, SXY) at each node is rotated "
        "into each ply's material coordinate system.  Tsai-Wu and Max Stress "
        "failure indices are evaluated for every ply at every node.  The envelope "
        "(worst ply) is reported.  The stress sign is taken from the dominant "
        "contributing mode (dominant-mode sign convention for SRSS)."
    )
    doc.add_paragraph(
        "Limitations: The analysis uses homogenised orthotropic properties for "
        "SOLID187 elements.  True ply-level through-thickness stress gradients "
        "are not captured.  This is a conservative preliminary assessment "
        "suitable for design-phase qualification screening."
    )

    _add_heading(doc, "7.2 Failure Results by Environment", level=2)

    # Master failure table
    fail_rows = []
    for prof in profiles:
        pid = prof["id"]
        for axis in ["X", "Y", "Z"]:
            key = (pid, axis)
            fr = failure_data.get(key)
            if fr:
                fail_rows.append([
                    prof["requirement_id"],
                    axis,
                    f"{fr.max_tw_index:.6f}",
                    f"{fr.max_ms_index:.6f}",
                    f"{fr.min_fos_tw:.1f}",
                    f"{fr.min_fos_ms:.1f}",
                    str(fr.critical_ply_tw[np.argmax(fr.tsai_wu_index)]),
                    _pass_fail_text(fr.min_fos_tw >= required_fos),
                ])

    _add_table(doc,
               ["Req ID", "Axis", "Max TW", "Max MS", "FoS_TW", "FoS_MS",
                "Crit Ply", "Result"],
               fail_rows)

    doc.add_paragraph("")
    _add_figure(doc, chart_paths.get("failure_bar"),
                "Figure 7-1: Composite failure indices by environment and axis")

    doc.add_paragraph("")
    _add_figure(doc, chart_paths.get("fos"),
                "Figure 7-2: Factor of Safety -- all environments")

    # Worst-case contour
    if contour_paths.get("tsai_wu"):
        doc.add_paragraph("")
        _add_figure(doc, contour_paths["tsai_wu"],
                    f"Figure 7-3: Tsai-Wu failure index contour -- worst case "
                    f"({worst_pid}, {worst_axis}-axis)")

    if contour_paths.get("fos_tw"):
        doc.add_paragraph("")
        _add_figure(doc, contour_paths["fos_tw"],
                    f"Figure 7-4: Tsai-Wu FoS contour -- worst case "
                    f"({worst_pid}, {worst_axis}-axis)")

    # Core failure (if data available)
    if core_failure_data:
        _add_heading(doc, "7.3 Core Failure Assessment", level=2)
        core_rows = []
        for (pid, axis), cf in core_failure_data.items():
            core_rows.append([
                pid, axis,
                f"{cf['max_index']:.6f}",
                f"{cf['min_fos']:.1f}",
                _pass_fail_text(cf["pass"]),
            ])
        _add_table(doc,
                   ["Environment", "Axis", "Max Index", "Min FoS", "Result"],
                   core_rows)

    doc.add_page_break()

    # ======================================================================
    # 8. REQUIREMENTS COMPLIANCE MATRIX
    # ======================================================================
    _add_heading(doc, "8. Requirements Compliance Matrix")

    comp_rows = []
    for prof in profiles:
        pid = prof["id"]
        # Find worst axis for this environment
        worst_tw = 0
        worst_fos = 999
        all_pass = True
        for axis in ["X", "Y", "Z"]:
            key = (pid, axis)
            fr = failure_data.get(key)
            if fr:
                worst_tw = max(worst_tw, fr.max_tw_index)
                worst_fos = min(worst_fos, fr.min_fos_tw)
                if fr.min_fos_tw < required_fos:
                    all_pass = False

        comp_rows.append([
            prof["requirement_id"],
            prof["name"],
            f"{worst_tw:.6f}",
            f"{worst_fos:.1f}",
            _pass_fail_text(all_pass),
            f"Section 6.{profiles.index(prof)+1}",
        ])

    _add_table(doc,
               ["Req ID", "Environment", "Worst TW Index", "Min FoS",
                "Result", "Reference"],
               comp_rows)

    doc.add_paragraph("")
    overall_pass = all(r[-2] == "PASS" for r in comp_rows) if comp_rows else False
    p = doc.add_paragraph()
    run = p.add_run(
        f"Overall Qualification: {'PASS' if overall_pass else 'FAIL'} "
        f"(Minimum FoS: {min_fos_tw:.1f} Tsai-Wu, {min_fos_ms:.1f} Max Stress)"
    )
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PASS_COLOR if overall_pass else FAIL_COLOR

    doc.add_page_break()

    # ======================================================================
    # 9. CONCLUSIONS
    # ======================================================================
    _add_heading(doc, "9. Conclusions")

    doc.add_paragraph(
        f"A random vibration (PSD) analysis was performed on the {config.part_name} "
        f"composite structure per MIL-STD-810H Method 514.8.  {len(profiles)} "
        f"vibration environments were evaluated with base excitation along all "
        f"three orthogonal axes."
    )

    _add_heading(doc, "9.1 Structural Adequacy", level=2)
    if overall_pass:
        doc.add_paragraph(
            f"The structure meets all MIL-STD-810H vibration requirements with "
            f"a minimum Tsai-Wu factor of safety of {min_fos_tw:.1f} "
            f"(required: {required_fos:.1f}).  "
            f"No composite failure is predicted under any tested environment."
        )
    else:
        doc.add_paragraph(
            f"The structure DOES NOT meet all requirements.  "
            f"Minimum Tsai-Wu FoS = {min_fos_tw:.1f} "
            f"(required: {required_fos:.1f}).  "
            f"Design modifications are recommended."
        )

    _add_heading(doc, "9.2 Dominant Modes", level=2)
    doc.add_paragraph(
        f"{modal_data.n_modes} modes extracted in "
        f"{modal_data.nat_freqs[0]:.1f}-{modal_data.nat_freqs[-1]:.1f} Hz range.  "
        f"Fundamental frequency: {modal_data.nat_freqs[0]:.1f} Hz."
    )

    _add_heading(doc, "9.3 Design Margins", level=2)
    doc.add_paragraph(
        f"Minimum Tsai-Wu FoS: {min_fos_tw:.1f} "
        f"(margin of {(min_fos_tw/required_fos - 1)*100:.0f}% above required).  "
        f"Minimum Max Stress FoS: {min_fos_ms:.1f}."
    )

    # Summary table
    _add_heading(doc, "9.4 Analysis Summary", level=2)
    summary = [
        ("Part", config.part_name),
        ("Geometry", os.path.basename(config.geometry_file)),
        ("Element Type", "SOLID187 (10-node tet)"),
        ("Mesh", f"{mi['n_elements']} elements, {mi['n_nodes']} nodes"),
        ("Material", config.material_name),
        ("Composite Layup", f"[{ls['stacking_sequence']}]"),
        ("Laminate Thickness", f"{ls['total_thickness_mm']:.3f} mm"),
        ("Modes Extracted", str(modal_data.n_modes)),
        ("Frequency Range", f"{modal_data.nat_freqs[0]:.1f}-{modal_data.nat_freqs[-1]:.1f} Hz"),
        ("Damping", f"{config.damping_ratio*100:.0f}% critical"),
        ("Environments Tested", str(len(profiles))),
        ("Required FoS", f"{required_fos:.1f}"),
        ("Min FoS (Tsai-Wu)", f"{min_fos_tw:.1f}"),
        ("Min FoS (Max Stress)", f"{min_fos_ms:.1f}"),
        ("Overall Result", verdict),
    ]
    _add_table(doc, ["Parameter", "Value"], summary)

    # ======================================================================
    # Save
    # ======================================================================
    output_path = os.path.join(output_dir, "MIL_STD_810H_PSD_Report.docx")
    doc.save(output_path)
    print(f"  Report saved: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _find_worst_case(failure_data):
    """Find the (profile_id, axis) with the highest Tsai-Wu index."""
    worst_pid, worst_axis, worst_fr = None, None, None
    worst_val = -1
    for (pid, axis), fr in failure_data.items():
        if fr.max_tw_index > worst_val:
            worst_val = fr.max_tw_index
            worst_pid = pid
            worst_axis = axis
            worst_fr = fr
    return worst_pid, worst_axis, worst_fr


def _worst_axis_for_env(pid, env_results, failure_data):
    """Find worst axis for a given environment."""
    worst_axis = None
    worst_val = -1
    for axis in ["X", "Y", "Z"]:
        key = (pid, axis)
        fr = failure_data.get(key)
        if fr and fr.max_tw_index > worst_val:
            worst_val = fr.max_tw_index
            worst_axis = axis
    return worst_axis
